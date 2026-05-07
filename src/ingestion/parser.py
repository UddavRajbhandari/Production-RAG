"""
Document Parser Module
Responsible for extracting raw content from various file formats (PDF, DOCX, XLSX).
Follows a modular architecture with type-specific handlers.
"""

import os
from abc import ABC, abstractmethod
from typing import Any

import docx
import openpyxl

try:
    import fitz  # PyMuPDF legacy import path
except ImportError:
    import pymupdf as fitz


class BaseHandler(ABC):
    """Abstract base class for format-specific document handlers."""

    @abstractmethod
    def parse(self, file_path: str) -> list[dict[str, Any]]:
        """Parses a file and returns a list of raw content blocks."""
        pass


class PDFHandler(BaseHandler):
    """Handler for PDF files using PyMuPDF (fitz)."""

    def parse(self, file_path: str) -> list[dict[str, Any]]:
        """
        Extracts text content on a per-page basis.
        Returns blocks containing page text and metadata.
        """
        blocks: list[dict[str, Any]] = []
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc):
            text = page.get_text().strip()
            blocks.append(
                {
                    "type": "page",
                    "content": text,
                    "metadata": {"page_number": page_num + 1, "char_count": len(text)},
                }
            )
        doc.close()
        return blocks


class DOCXHandler(BaseHandler):
    """Handler for Word documents using python-docx."""

    def parse(self, file_path: str) -> list[dict[str, Any]]:
        """
        Traverses paragraphs and tables in a Word document.
        Preserves paragraph styles and basic table structure.
        """
        blocks: list[dict[str, Any]] = []
        doc = docx.Document(file_path)
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                blocks.append(
                    {
                        "type": "paragraph",
                        "content": para.text.strip(),
                        "metadata": {
                            "paragraph_index": i,
                            "style": para.style.name if para.style else "Normal",
                        },
                    }
                )
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            blocks.append(
                {"type": "table", "content": table_data, "metadata": {"table_index": i}}
            )
        return blocks


class XLSXHandler(BaseHandler):
    """Handler for Excel files using openpyxl."""

    def parse(self, file_path: str) -> list[dict[str, Any]]:
        """
        Iterates through sheets and rows in an Excel workbook.
        Each sheet is treated as a distinct unit.
        """
        blocks: list[dict[str, Any]] = []
        wb = openpyxl.load_workbook(file_path, data_only=True)
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            data = []
            for row in sheet.iter_rows(values_only=True):
                if any(row):  # Skip empty rows
                    data.append(list(row))
            blocks.append(
                {
                    "type": "sheet",
                    "content": data,
                    "metadata": {"sheet_name": sheet_name},
                }
            )
        return blocks


class DocumentParser:
    """Entry point for parsing files based on their extension."""

    def __init__(self) -> None:
        """Initializes supported handlers."""
        self.handlers: dict[str, BaseHandler] = {
            ".pdf": PDFHandler(),
            ".docx": DOCXHandler(),
            ".xlsx": XLSXHandler(),
        }

    def parse(self, file_path: str) -> list[dict[str, Any]]:
        """
        Determines the correct handler and executes the parse operation.
        Raises ValueError if extension is unsupported.
        """
        ext = os.path.splitext(file_path)[1].lower()
        handler = self.handlers.get(ext)
        if not handler:
            raise ValueError(f"No handler for file extension: {ext}")

        return handler.parse(file_path)
