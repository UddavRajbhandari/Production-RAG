"""
Expanded Corpus Audit Script
Randomly samples 20+ files from the raw directory for broader validation.
Includes checks for zero-text pages and complex nested structures.
"""

import os
import random

import docx
import fitz  # PyMuPDF


def audit_pdf(file_path: str) -> None:
    """Deep audit of first 3 and last page of a PDF."""
    print(f"\n--- Auditing PDF: {os.path.basename(file_path)} ---")
    try:
        doc = fitz.open(file_path)
        print(f"Pages: {len(doc)}")
        check_pages = list(range(min(3, len(doc))))
        if len(doc) > 3:
            check_pages.append(len(doc) - 1)

        for i in check_pages:
            page = doc[i]
            text = page.get_text().strip()
            tables = page.find_tables()
            print(
                f"Page {i + 1}: {len(text)} chars, "
                f"{len(tables.tables)} potential tables detected"
            )
            if len(text) < 50:
                print(
                    f"  [!] Note: Very low text on page {i + 1} (Cover/Image/Separator)"
                )
        doc.close()
    except Exception as e:
        print(f"  [!] CRITICAL: Error opening PDF: {e}")


def audit_docx(file_path: str) -> None:
    """Audit of DOCX structure and table counts."""
    print(f"\n--- Auditing DOCX: {os.path.basename(file_path)} ---")
    try:
        doc = docx.Document(file_path)
        print(f"Paragraphs: {len(doc.paragraphs)}")
        print(f"Tables: {len(doc.tables)}")
    except Exception as e:
        print(f"  [!] CRITICAL: Error opening DOCX: {e}")


def get_all_files(directory: str, extensions: list[str]) -> list[str]:
    """Recursive file discovery based on extensions."""
    found_files = []
    for root, _, files in os.walk(directory):
        for file in files:
            if any(file.lower().endswith(ext) for ext in extensions):
                found_files.append(os.path.join(root, file))
    return found_files


# Define search parameters
raw_dir = "data/raw"
exts = [".pdf", ".docx"]

if __name__ == "__main__":
    all_files = get_all_files(raw_dir, exts)
    already_audited = [
        "Access-to-Information-2023-annual-report.pdf",
        "2605.02520v1.pdf",
        "World-Bank-Access-to-Information-FY22-annual-report.pdf",
        "f3b45293bf4f8d691cb5330a2d17974b0fdcbff7a1d6dad57f57b8b9b11fbf3f.docx",
        "2605.02661v1.pdf",
    ]
    files_to_audit = [
        f for f in all_files if os.path.basename(f) not in already_audited
    ]
    sample_size = min(20, len(files_to_audit))
    selected_files = random.sample(files_to_audit, sample_size)

    print(f"Found {len(all_files)} files total. Auditing {sample_size} new files...")

    for file in selected_files:
        if file.lower().endswith(".pdf"):
            audit_pdf(file)
        elif file.lower().endswith(".docx"):
            audit_docx(file)
