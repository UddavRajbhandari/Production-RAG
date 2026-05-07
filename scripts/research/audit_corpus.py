"""
Corpus Audit Script
Performs a sanity check on a small sample of PDF and DOCX files.
Identifies initial layout, text yield, and table presence.
"""

import os

import docx
import fitz  # PyMuPDF


def audit_pdf(file_path: str) -> None:
    """Logs page count, text length, and table detection for PDFs."""
    print(f"\n--- Auditing PDF: {os.path.basename(file_path)} ---")
    try:
        doc = fitz.open(file_path)
        print(f"Pages: {len(doc)}")
        for i in range(min(3, len(doc))):
            page = doc[i]
            text = page.get_text()
            tables = page.find_tables()
            print(
                f"Page {i + 1}: {len(text)} chars, "
                f"{len(tables.tables)} potential tables detected"
            )
            if len(text) < 100:
                print(
                    f"  [!] Warning: Very low text on page {i + 1} "
                    "(Possible image-only?)"
                )
        doc.close()
    except Exception as e:
        print(f"  [!] Error auditing PDF: {e}")


def audit_docx(file_path: str) -> None:
    """Logs paragraph and table counts for DOCX files."""
    print(f"\n--- Auditing DOCX: {os.path.basename(file_path)} ---")
    try:
        doc = docx.Document(file_path)
        print(f"Paragraphs: {len(doc.paragraphs)}")
        print(f"Tables: {len(doc.tables)}")
    except Exception as e:
        print(f"  [!] Error auditing DOCX: {e}")


# Sample files for audit based on directory listing
samples: list[str] = [
    "data/raw/pdf/Access-to-Information-2023-annual-report.pdf",
    "data/raw/pdf/2605.02520v1.pdf",
    "data/raw/pdf/World-Bank-Access-to-Information-FY22-annual-report.pdf",
    "data/raw/docx/f3b45293bf4f8d691cb5330a2d17974b0fdcbff7a1d6dad57f57b8b9b11fbf3f.docx",
    "data/raw/pdf/2605.02661v1.pdf",
]

if __name__ == "__main__":
    for sample in samples:
        if os.path.exists(sample):
            if sample.endswith(".pdf"):
                audit_pdf(sample)
            elif sample.endswith(".docx"):
                audit_docx(sample)
        else:
            print(f"File not found: {sample}")
